

import streamlit as st
from docx import Document
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import string
import pandas as pd
import pickle
import sklearn
from sklearn.metrics import mean_squared_error, r2_score
from sklearn.preprocessing import MinMaxScaler
from sklearn.preprocessing import StandardScaler
import joblib
import numpy as np
import docx2txt
import matplotlib.pyplot as plt
import re
from io import BytesIO



# Download NLTK stopwords (if not already installed)
nltk.download('punkt_tab')
nltk.download('stopwords')


# Set page config
st.set_page_config(page_title="Streamlit Sidebar Navbar", layout="wide")

# Custom CSS for sidebar styling
st.markdown(
    """
    <style>
        /* Style sidebar */
        [data-testid="stSidebar"] {
            background-color: #001F3F !important;  /* Navy Blue */
        }

        /* Center buttons and make them equal width */
        .sidebar-button {
            display: flex;
            justify-content: center;
            align-items: center;
            text-align: center;
            width: 100% !important;
            color: white !important;
            background-color: transparent;
            border: 2px solid white;
            padding: 10px;
            margin: 10px 0;
            font-size: 18px;
            border-radius: 8px;
            cursor: pointer;
        }

        /* Hover effect */
        .sidebar-button:hover {
            background-color: white;
            color: #001F3F !important;
        }

        /* Center align the buttons inside the sidebar */
        [data-testid="stSidebarContent"] {
            display: flex;
            flex-direction: column;
            align-items: center;
        }
    </style>
    """,
    unsafe_allow_html=True,
)
st.sidebar.title("🔹 SKILL BRIDGE")

# Initialize session state for page switching
if "page" not in st.session_state:
    st.session_state.page = "Home"

# Sidebar buttons with centered alignment
if st.sidebar.button("🏠 Check your resume score now", key="home"):
    st.session_state.page = "Home"

if st.sidebar.button("Get Recommendations ", key="about"):
    st.session_state.page = "Recommendation"

if st.sidebar.button("Salary prediction", key="contact"):
    st.session_state.page = "Salary_Prediction"

if st.sidebar.button("Generate Resume", key="resume"):
    st.session_state.page="Generate"

# Display the selected page content
st.title(f"{st.session_state.page}")

if st.session_state.page == "Home":

    nltk.download('punkt_tab')
    nltk.download('stopwords')

    # Skill keywords for web development and data science analysis
    skills_keywords = set([
        # Web Development
        'html', 'css', 'javascript', 'react', 'angular', 'vue', 'node.js', 'python', 'django',
        'flask', 'php', 'sql', 'mongodb', 'rest', 'api', 'git', 'responsive design', 'ui', 'ux',

        # Data Science & Analysis
        'r', 'pandas', 'numpy', 'matplotlib', 'seaborn', 'scikit-learn',
        'tensorflow', 'pytorch', 'machine learning', 'deep learning', 'data mining',
        'data analysis', 'data visualization', 'statistical modeling', 'nlp', 'big data',
        'hadoop', 'spark', 'tableau', 'power bi'
    ])

    experience_patterns = [
        r"\b(\d+)\+? years?\b",
        r"\b(\d+)\s+years?\s+of\s+experience\b",
    ]

    cloud_api_patterns = [
        r"\bAWS\b", r"\bAzure\b", r"\bGCP\b",
        r"\bREST API\b", r"\bAPI development\b", r"\bRESTful APIs\b",
        r"\bmicroservices architecture\b", r"\bcontainerization \(Docker\)\b",
        r"\bCI\/CD\b", r"\bGitHub Actions\b"
    ]

    necessary_skill_patterns = [
        r"must have\s+(\w+)",
    ]

    github_patterns = [
        r"\bgithub\.com\/[^\s]+\b",
        r"\bgit\b",
        r"\bversion control tools like Git\b"
    ]

    pipeline_patterns = [
        r"\bCI\/CD\b",
        r"\bpipeline\b",
    ]

    jd_mentions_patterns = [
        r"\b(JD|Job Description)\b mentions\s*(.*)",
    ]


    def extract_text_from_docx(uploaded_file):
        return docx2txt.process(uploaded_file)


    def extract_skills(text):
        stop_words = set(stopwords.words('english'))
        tokens = word_tokenize(text)
        filtered_tokens = [w for w in tokens if w.lower() not in stop_words and w.isalnum()]
        skills = [token.lower() for token in filtered_tokens if token.lower() in skills_keywords]

        experience = [match for pattern in experience_patterns for match in re.findall(pattern, text, re.IGNORECASE)]
        cloud_apis = [match for pattern in cloud_api_patterns for match in re.findall(pattern, text, re.IGNORECASE)]
        necessary_skills = [match for pattern in necessary_skill_patterns for match in
                            re.findall(pattern, text, re.IGNORECASE)]
        github_info = [match for pattern in github_patterns for match in re.findall(pattern, text, re.IGNORECASE)]
        pipeline_info = [match for pattern in pipeline_patterns for match in re.findall(pattern, text, re.IGNORECASE)]
        jd_mentions = [match for pattern in jd_mentions_patterns for match in re.findall(pattern, text, re.IGNORECASE)]

        return skills, experience, cloud_apis, necessary_skills, github_info, pipeline_info, jd_mentions


    st.title("Resume vs Job Description Skill Match Analyzer")

    resume_file = st.file_uploader("Upload Resume (.docx)", type=["docx"])
    jd_file = st.file_uploader("Upload Job Description (.docx)", type=["docx"])

    if resume_file and jd_file:
        resume_text = extract_text_from_docx(resume_file)
        jd_text = extract_text_from_docx(jd_file)

        resume_data = extract_skills(resume_text)
        jd_data = extract_skills(jd_text)

        resume_skills, resume_experience, resume_cloud_apis, resume_necessary_skills, resume_github_info, resume_pipeline_info, resume_jd_mentions = resume_data
        jd_skills, jd_experience, jd_cloud_apis, jd_necessary_skills, jd_github_info, jd_pipeline_info, jd_jd_mentions = jd_data

        common_skills = set(resume_skills) & set(jd_skills)
        similarity_score = len(common_skills) / len(jd_skills) * 100 if jd_skills else 0

        labels = ["Matching Skills", "Missing Skills"]
        sizes = [similarity_score, 100 - similarity_score]

        fig, ax = plt.subplots()
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90)
        ax.axis('equal')
        st.pyplot(fig)

        st.subheader("Extracted Details from Resume")
        if resume_skills: st.write("**Skills:**", resume_skills)
        if resume_experience: st.write("**Experience:**", resume_experience)
        if resume_cloud_apis: st.write("**Cloud APIs:**", resume_cloud_apis)
        if resume_necessary_skills: st.write("**Necessary Skills:**", resume_necessary_skills)
        if resume_github_info: st.write("**GitHub Info:**", resume_github_info)
        if resume_pipeline_info: st.write("**Pipeline Info:**", resume_pipeline_info)
        if resume_jd_mentions: st.write("**JD Mentions:**", resume_jd_mentions)

        st.subheader("Extracted Details from JD")
        if jd_skills: st.write("**Skills:**", jd_skills)
        if jd_experience: st.write("**Experience:**", jd_experience)
        if jd_cloud_apis: st.write("**Cloud APIs:**", jd_cloud_apis)
        if jd_necessary_skills: st.write("**Necessary Skills:**", jd_necessary_skills)
        if jd_github_info: st.write("**GitHub Info:**", jd_github_info)
        if jd_pipeline_info: st.write("**Pipeline Info:**", jd_pipeline_info)
        if jd_jd_mentions: st.write("**JD Mentions:**", jd_jd_mentions)

        st.success(f"Skill Match Score: {similarity_score:.2f}%")


# Print first 200 characters to check


elif st.session_state.page == "Recommendation":
    st.write("This is the About page. Here, you can describe your project.")

    with open("pathname/recommendationmodel", 'rb') as f:
        data, cv, similarity = pickle.load(f)


    def content_recomm(course):
        try:
            course_index = data[data['course_title'] == course].index[0]
            sim = similarity[course_index]
            course_list = sorted(enumerate(sim), reverse=True, key=lambda x: x[1])[1:6]
            recommendations = [data.iloc[i[0]]['course_title'] for i in course_list]
            return recommendations
        except IndexError:
            st.error("Course not found in the dataset.")
            return []


    # Streamlit app code
    st.title("Course Recommender")
    course_list = data['course_title'].unique().tolist()
    selected_course = st.selectbox("Select a course:", course_list)

    if st.button("Recommend"):
        recommendations = content_recomm(selected_course)
        st.write("Recommended Courses:")
        for recommendation in recommendations:
            st.write(recommendation)



elif st.session_state.page == "Salary_Prediction":




    # --- Streamlit UI ---
    st.title("💼 Engineering Graduate Salary Predictor")
    model = pickle.load(open("path/salarymodel.pkl", 'rb'))
    import joblib
    std_scaler = joblib.load("path/std_scaler.pkl")
    label_encoder = joblib.load("path/label_encoder.pkl")
    st.markdown("Enter your details below to predict your estimated salary:")
    st.title("Employee Salary Prediction")

    # Input fields for features
    age = st.number_input("Age", min_value=18, max_value=65, value=30)
    gender = st.selectbox("Gender", ["Male", "Female"])
    degree = st.selectbox("Degree", ["Bachelors", "Masters", "PhD"])
    job_title = st.selectbox("Job Title", ["Data Scientist", "Software Engineer", ...])  # Add all job titles here
    experience_years = st.number_input("Experience Years", min_value=0, max_value=30, value=5)

    # Preprocess the input data
    # Create a DataFrame from the input values
    input_data = pd.DataFrame({
        "Age": [age],
        "Gender": [gender],
        "Degree": [degree],
        "job_Title": [job_title],
        "Experience_years": [experience_years]
    })

# Encode categorical features using LabelEncoder
    input_data["Gender_Encode"] = label_encoder.fit_transform(input_data["Gender"])
    input_data["Degree_Encode"] = label_encoder.fit_transform(input_data["Degree"])
    input_data["job_Title_Encode"] = label_encoder.fit_transform(input_data["job_Title"])

    # Scale numerical features using StandardScaler
    # Reshape the input data for scaling
    age_scaled = std_scaler.transform([[age]])[0][0]  # Reshape age
    experience_years_scaled = std_scaler.transform([[experience_years]])[0][0]  # Reshape Experience_years
    input_data["Age_scaled"] = age_scaled
    input_data["Experience_years_scaled"] = experience_years_scaled

    # Select the features for prediction
    features = ["Age_scaled", "Experience_years_scaled", "Gender_Encode", "Degree_Encode", "job_Title_Encode"]
    input_data_for_prediction = input_data[features]


    # Make prediction using the loaded model
    prediction = model.predict(input_data_for_prediction)

    # Display the prediction
    st.write(f"Predicted Salary: ${prediction[0]:.2f}")
    st.write("This is the Contact page. You can put your contact details here.")


elif st.session_state.page == "Generate":
    st.write("This is the Contact page. You can put your contact details here.")
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO


    def generate_resume():
        document = Document()

        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # Name & Contact
        name_heading = document.add_heading(name, 0)
        name_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para = document.add_paragraph(f"Email: {email} | Phone: {phone} | LinkedIn: {linkedin}")
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Introduction
        document.add_heading("Introduction", level=1)
        document.add_paragraph(introduction)

        # Education Table
        document.add_heading("Education", level=1)
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Qualification'
        hdr_cells[1].text = 'Board/University'
        hdr_cells[2].text = 'Percentage/CGPA'

        edu_rows = [
            ("Class 10", tenth_board, tenth_marks),
            ("Class 12", twelfth_board, twelfth_marks),
            ("College", university, college_cgpa)
        ]
        for row in edu_rows:
            cells = table.add_row().cells
            cells[0].text = row[0]
            cells[1].text = row[1]
            cells[2].text = row[2]

        # Skills
        document.add_heading("Skills", level=1)
        document.add_paragraph("• Technical Skills: " + " | ".join([s.strip() for s in tech_skills]))
        document.add_paragraph("• Soft Skills: " + " | ".join([s.strip() for s in soft_skills]))

        # Internships
        if internships:
            document.add_heading("Internship Experience", level=1)
            for internship in internships:
                document.add_paragraph(internship, style='List Bullet')
                document.add_paragraph("• [Explain your role, contributions, technologies used, and outcomes]")

        # Projects
        if projects:
            document.add_heading("Notable Projects", level=1)
            for project in projects:
                document.add_paragraph(project, style='List Bullet')
                document.add_paragraph("• [Describe functionality, tools used, and end result]")

        # Achievements
        if achievements:
            document.add_heading("Achievements", level=1)
            for achievement in achievements:
                document.add_paragraph(achievement, style='List Bullet')

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer


    # Streamlit UI
    st.title("📄 Resume Generator")

    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone")
    linkedin = st.text_input("LinkedIn Profile URL")
    introduction = st.text_area("Brief Introduction")

    st.subheader("🎓 Education")
    tenth_board = st.text_input("10th Board")
    tenth_marks = st.text_input("10th Marks / Percentage")
    twelfth_board = st.text_input("12th Board")
    twelfth_marks = st.text_input("12th Marks / Percentage")
    university = st.text_input("University Name")
    college_cgpa = st.text_input("College CGPA")

    st.subheader("🛠 Skills")
    tech_skills = st.text_input("Technical Skills (comma separated)").split(",")
    soft_skills = st.text_input("Soft Skills (comma separated)").split(",")

    st.subheader("💼 Internships")
    internship_count = st.number_input("Number of Internships", min_value=0, step=1)
    internships = [st.text_input(f"Internship {i + 1}") for i in range(internship_count)]

    st.subheader("🚀 Projects")
    project_count = st.number_input("Number of Projects", min_value=0, step=1)
    projects = [st.text_input(f"Project {i + 1}") for i in range(project_count)]

    st.subheader("🏆 Achievements")
    achievement_count = st.number_input("Number of Achievements", min_value=0, step=1)
    achievements = [st.text_input(f"Achievement {i + 1}") for i in range(achievement_count)]

    # Generate & Download
    if st.button("Generate Resume"):
        resume_file = generate_resume()
        st.download_button(
            label="📥 Download Resume",
            data=resume_file,
            file_name="resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    st.write("This is the Contact page. You can put your contact details here.")
